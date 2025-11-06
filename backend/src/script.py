import sys
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT

# --- Helper Function for All-Caps Section Headings ---
def add_section_heading(doc, text, align=WD_PARAGRAPH_ALIGNMENT.LEFT, space_before=12, space_after=6):
    p = doc.add_paragraph()
    p.alignment = align
    run = p.add_run(text)
    run.bold = True
    run.font.all_caps = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)

# --- Helper Function for Key: Value Pairs ---
def add_indented_detail(doc, key, value):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.5)
    p.add_run(f'{key}:').bold = True
    p.add_run(f' {value.strip()}')

# --- 1. Create Document & Set Default Styles ---
doc = Document()

# Set default font to Times New Roman, 12pt, and 1.5 line spacing
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)
style.paragraph_format.line_spacing = 1.5

# Define Heading 1 Style (for top titles)
h1_style = doc.styles.add_style('MainHeading', 1)
h1_style.font.name = 'Times New Roman'
h1_style.font.size = Pt(14)
h1_style.font.bold = True
h1_style.font.all_caps = True
h1_style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
h1_style.paragraph_format.space_after = Pt(12)
h1_style.paragraph_format.space_before = Pt(0)

# --- 2. Add Content ---

# Top Headings
doc.add_paragraph('Annexure3b- Complete Filing', style='MainHeading')
doc.add_paragraph('Invention Disclosure Form', style='MainHeading')
doc.add_paragraph('Detailed Description of the Proposed Invention')

# Section 1: TITLE
p_title = doc.add_paragraph()
p_title.add_run('1. Title: ').bold = True
p_title.add_run('AI-Powered Real-Time Public Transport Tracking and Crowd Density Management System for Small Cities')
doc.add_paragraph('This invention presents an innovative web-based platform designed to transform public transportation in small urban centers by delivering real-time vehicle tracking, precise arrival time predictions, and crowd density management. Specifically crafted for regions with limited technological infrastructure, the system ensures accessibility for all commuters, including those without smartphones, through a user-friendly website, SMS updates, and interactive voice response (IVR) services. By leveraging artificial intelligence and data aggregation, it provides reliable transit information and empowers city authorities with a comprehensive dashboard to optimize operations, enhance efficiency, and promote sustainable urban mobility.')
doc.add_paragraph('Public transport in small cities often suffers from inconsistent schedules, overcrowded vehicles, and a lack of accessible real-time information, leading to commuter frustration and inefficient fleet management. This AI-powered platform addresses these challenges by integrating live data from transport operators, using machine learning to predict arrival times, and estimating crowd levels through user inputs and schedule patterns. Its software-only design ensures functionality in low-connectivity environments, making it ideal for semi-urban areas. The system’s inclusive approach ensures equitable access for all residents, fostering community engagement and trust.')
doc.add_paragraph('By enabling better trip planning and operational efficiency, the platform reduces wait times, balances passenger loads, and encourages public transport use over private vehicles, contributing to reduced traffic congestion and carbon emissions. Its cost-effective, hardware-free approach makes it feasible for small municipalities with limited budgets. The system’s scalability supports future integrations, such as digital payments or smart city frameworks, positioning it as a forward-thinking solution for modernizing public transportation in underserved urban communities.')

# Section 2: INTERNAL INVENTOR(S)/ STUDENT(S)
p_inventors = doc.add_paragraph()
p_inventors.paragraph_format.space_before = Pt(12)
p_inventors.add_run('2. Internal Inventor(s)/ Student(s): ').bold = True
p_inventors.add_run('All fields in this section are mandatory.')
doc.add_paragraph('This section documents the details of internal inventors or students who contributed to the development of this AI-powered public transport system. The table provided below captures essential information to ensure proper recognition and compliance with patent filing requirements, reflecting the collaborative effort behind this innovation.')
doc.add_paragraph('The invention was developed through a multidisciplinary team effort, combining expertise in software engineering, artificial intelligence, and urban mobility solutions. Each contributor played a critical role in addressing the unique challenges of small-city transportation, ensuring the platform meets the needs of diverse users and authorities. The structured format below allows for clear documentation of their involvement, with flexibility to accommodate additional contributors if needed.')

internal_table = doc.add_table(rows=7, cols=6)
internal_table.style = 'Table Grid'
internal_table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr_cells = internal_table.rows[0].cells
hdr_texts = ['Full Name', 'Mobile', 'Email', 'UID/Reg No', 'Address', 'Signature']

for i, text in enumerate(hdr_texts):
    p = hdr_cells[i].paragraphs[0]
    p.add_run(text).bold = True
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

for i in range(1, 7):
    row_cells = internal_table.rows[i].cells
    row_cells[0].text = f'[Inventor {i} Name]'
    row_cells[1].text = '[Mobile]'
    row_cells[2].text = '[Email]'
    row_cells[3].text = '[UID/Reg No]'
    row_cells[4].text = '[Address]'
    row_cells[5].text = '[Signature]'

doc.add_paragraph('For additional inventors, please add rows as needed.')

# External Inventors
p_ext_inventors = doc.add_paragraph()
p_ext_inventors.add_run('External Inventor(s): (Inventors Not Affiliated with LPU)').bold = True
doc.add_paragraph('This section is reserved for documenting any external inventors who are not affiliated with Lovely Professional University but have contributed to the invention. A No Objection Certificate (NOC) from their affiliated institute, university, industry, or laboratory is mandatory to clarify intellectual property rights and ensure smooth filing processes.')
doc.add_paragraph('The inclusion of external contributors, if applicable, enhances the collaborative scope of this project, bringing diverse perspectives to the development process. The NOC requirement ensures all parties agree on ownership and commercialization terms, preventing potential disputes. The table below provides a structured format for capturing their details, with the NOC format attached for reference.')

external_table = doc.add_table(rows=2, cols=6)
external_table.style = 'Table Grid'
external_table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr_cells_ext = external_table.rows[0].cells

for i, text in enumerate(hdr_texts):
    p = hdr_cells_ext[i].paragraphs[0]
    p.add_run(text).bold = True
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

row_cells = external_table.rows[1].cells
row_cells[0].text = '[External Inventor Name]'
row_cells[1].text = '[Mobile]'
row_cells[2].text = '[Email]'
row_cells[3].text = '[UID/Reg No]'
row_cells[4].text = '[Address]'
row_cells[5].text = '[Signature]'
doc.add_paragraph()

# Section 3: DESCRIPTION OF THE INVENTION
p_desc = doc.add_paragraph()
p_desc.paragraph_format.space_before = Pt(12)
p_desc.add_run('3. Description of the Invention: ').bold = True
doc.add_paragraph('The AI-Powered Real-Time Public Transport Tracking and Crowd Density Management System is a transformative web-based platform designed to enhance public transportation in small urban centers. These areas often face challenges such as unreliable schedules, overcrowded buses, and limited access to real-time information, particularly for residents without smartphones. Unlike solutions tailored for large cities, which rely on high-speed internet and complex hardware, this system is entirely software-based, leveraging existing data sources to deliver accurate transit updates and crowd management insights.')
doc.add_paragraph('The platform aggregates data from transport operators, such as driver-reported locations or open transit feeds, and processes it through a cloud-based system powered by advanced AI algorithms. These algorithms analyze historical and real-time data, including traffic patterns and schedules, to predict vehicle arrival times with high precision. Crowd density is estimated using user-reported inputs or predictive models based on transit patterns, ensuring safer and more comfortable rides without requiring physical sensors. Commuters can access this information through an intuitive website, SMS updates, or IVR services, making the system inclusive for diverse populations, including rural and elderly users.')
doc.add_paragraph('For city authorities, the platform provides a robust management dashboard that visualizes fleet performance, passenger distribution, and operational trends, enabling data-driven decisions to optimize routes and reduce delays. The system’s lightweight design, using efficient data protocols, ensures reliable performance in low-connectivity environments, a critical feature for small cities. Security measures, including encrypted data transmission and compliance with privacy regulations, protect user information and foster trust.')
doc.add_paragraph('The development process involved extensive stakeholder engagement to identify pain points, followed by iterative design and testing to ensure usability and reliability. The platform’s technical framework combines advanced analytics with user-centric design, achieving strong performance in simulated environments. Its focus on inclusivity, sustainability, and cost-effectiveness positions it as a groundbreaking solution for small urban communities seeking to modernize their public transport systems.')

# Development Phases
doc.add_paragraph('Various phases are described below:')
phases = [
    ('1. Conceptualization and Ideation', 'The development began with a thorough exploration of the challenges faced by small-city commuters and transport operators. Through surveys and discussions, the team identified key issues such as unreliable schedules, overcrowding, and lack of accessible information. Collaborative brainstorming sessions, involving experts in software engineering, AI, and urban mobility, led to the conceptualization of a web-based platform that integrates real-time tracking, predictive analytics, and inclusive access channels. Initial wireframes for the website and dashboard were created, focusing on simplicity and adaptability to low-resource environments.'),
    ('2. Research and Feasibility Study', 'This phase involved a comprehensive review of existing public transport technologies and their limitations in small cities. The team analyzed data aggregation methods, predictive algorithms, and low-bandwidth communication protocols to ensure feasibility. Simulations were conducted to model transit patterns and test system performance in low-connectivity scenarios. Economic and regulatory analyses confirmed the platform’s viability, highlighting its potential to deliver cost-effective solutions while adhering to data privacy standards.'),
    ('3. Prototype Design and Integration', 'A functional prototype was developed, focusing on seamless data integration and user-friendly interfaces. The design prioritized intuitive navigation for commuters and actionable insights for authorities. Integration testing ensured smooth data flow from transit sources to the cloud-based system. User experience testing with diverse groups refined the website, SMS, and IVR interfaces, ensuring accessibility and ease of use across different demographics.'),
    ('4. AI-Driven Algorithm Development', 'Advanced machine learning models were developed to predict arrival times and crowd density, using historical and real-time transit data. These models were optimized for speed and accuracy, ensuring reliable performance in dynamic conditions. The algorithms were integrated into the platform via secure APIs, with fairness checks to prevent biases in predictions. This phase also included fine-tuning to handle low-bandwidth environments effectively.'),
    ('5. Testing and Validation', 'Rigorous testing was conducted in simulated small-city environments to validate the platform’s performance. Tests demonstrated high accuracy in arrival predictions and effective crowd density estimation, with minimal latency. User feedback refined the interface, enhancing accessibility and usability. Scalability tests confirmed the system’s ability to handle multiple data streams, ensuring readiness for real-world deployment.'),
    ('6. Regulatory Approval and Compliance', 'The platform underwent evaluation to ensure compliance with data privacy and telecommunications regulations. Detailed documentation, including risk assessments and technical reports, was prepared for regulatory bodies. Third-party audits verified the system’s security and reliability, aligning with industry standards. This phase ensured the platform’s eligibility for widespread adoption in public transport systems.'),
    ('7. Commercial Production and Launch', 'Plans for commercial rollout were developed, focusing on partnerships with municipal transport authorities. Pilot deployments were designed to demonstrate the platform’s impact in small cities. Stakeholder workshops and training programs were planned to ensure smooth adoption, with marketing efforts highlighting the system’s benefits for commuters and operators. The launch strategy emphasizes phased rollouts to maximize impact.'),
    ('8. Continuous Improvement and AI Learning', 'Post-launch, the platform collects real-time data to refine AI models, improving prediction accuracy and incorporating user feedback. Performance monitoring tools track key metrics, with automated alerts for anomalies. User forums enable iterative enhancements, ensuring the system remains relevant. Future updates could include integrations with advanced networks or smart city platforms, enhancing long-term utility.')
]

for phase_title, phase_desc in phases:
    p_phase_title = doc.add_paragraph()
    p_phase_title.add_run(phase_title).bold = True
    doc.add_paragraph(phase_desc)
    doc.add_paragraph()

# PROBLEM ADDRESSED BY THE INVENTION
add_section_heading(doc, 'Problem Addressed by the Invention:')
doc.add_paragraph('Public transportation in small cities is often hampered by unreliable schedules, overcrowded vehicles, and limited access to real-time information, particularly for residents without smartphones. These issues lead to prolonged wait times, uneven passenger distribution, and increased reliance on private vehicles, exacerbating traffic congestion and environmental concerns. City authorities rely on manual processes, which lack the flexibility to respond to dynamic transit demands.')
doc.add_paragraph('Existing digital solutions, designed for large metropolitan areas, require robust internet infrastructure and exclude non-digital users, making them impractical for small cities. These platforms often lack features for crowd density management without hardware, limiting their ability to ensure passenger safety and comfort. This invention addresses these challenges by providing a software-only platform that delivers real-time tracking, accurate arrival predictions, and crowd insights, accessible to all users.')
doc.add_paragraph('The platform’s inclusive design, with website, SMS, and IVR access, ensures that rural and elderly residents are not excluded. For authorities, it offers a dashboard to optimize operations, reducing inefficiencies and improving service quality. By promoting public transport use, the system contributes to reduced traffic and emissions, fostering sustainable urban development in small communities.')

# Problem / Solution Summary
p_problem = doc.add_paragraph()
p_problem.add_run('Problem: ').bold = True
p_problem.add_run('Small cities face unreliable public transport schedules, overcrowding, and limited information access for non-digital users, with manual management leading to inefficiencies.')
doc.add_paragraph('These challenges result in commuter frustration, underutilized vehicles, and environmental strain from increased private vehicle use. Existing tools, built for larger cities, fail to address connectivity constraints and the digital divide in smaller urban areas.')

p_solution = doc.add_paragraph()
p_solution.add_run('Solution: ').bold = True
p_solution.add_run('A web-based platform delivering real-time tracking, AI-predicted arrival times, and crowd density insights, accessible via website, SMS, or IVR, with a dashboard for authorities.')
doc.add_paragraph('In a typical small city with a modest bus fleet, manual scheduling leads to significant inefficiencies. This platform can reduce wait times and improve fleet utilization, as shown in early tests. Its inclusive access channels ensure all residents benefit, addressing critical gaps in existing transit solutions.')

# OBJECTIVE OF THE INVENTION
add_section_heading(doc, 'Objective of the Invention')
doc.add_paragraph('The primary objective of this invention is to create a web-based system that delivers real-time public transport information tailored for small cities, enabling commuters to access accurate and timely updates. The platform aims to provide precise arrival time predictions and crowd density insights using AI-driven analytics, enhancing trip planning and passenger safety.')
doc.add_paragraph('It seeks to ensure inclusivity by offering multiple access channels, including a website, SMS, and IVR, to accommodate users without smartphones. For transport authorities, the system provides a centralized dashboard to monitor fleet performance and make data-driven decisions, improving operational efficiency. The platform promotes environmental sustainability by encouraging public transport use, reducing traffic congestion and emissions.')
doc.add_paragraph('Designed for low-connectivity environments, the system incorporates robust security measures to protect user data. Its scalable architecture supports future integrations, such as digital payments or smart city connectivity, ensuring long-term relevance. Through rigorous testing and stakeholder collaboration, the platform aims to achieve high reliability and user satisfaction, facilitating widespread adoption.')

# C. STATE OF THE ART/ RESEARCH GAP/ NOVELTY
add_section_heading(doc, 'C. State of the Art/ Research Gap/ Novelty: How does the invention address the research gap?')
doc.add_paragraph('Current public transport solutions are primarily designed for large metropolitan areas with reliable internet and high smartphone penetration. These systems, such as popular transit apps, often fail to meet the needs of small cities, where connectivity is limited and many residents rely on basic phones. Existing patents focus on scheduling or basic tracking but lack features for crowd density management without hardware and do not prioritize inclusivity.')
doc.add_paragraph('This invention addresses these gaps by offering a software-only platform optimized for low-bandwidth environments. Its AI-driven approach to predicting arrival times and crowd levels, using user inputs and transit patterns, eliminates the need for costly sensors. The inclusion of SMS and IVR access ensures that all residents, including rural and elderly users, can benefit, bridging the digital divide.')
doc.add_paragraph('The platform’s management dashboard provides actionable insights, a feature often missing in existing tools. Its focus on small urban communities, combined with its emphasis on sustainability and cost-effectiveness, sets it apart from metropolitan-focused solutions. Research highlights connectivity and inclusivity as key challenges, which this system addresses through efficient protocols and multi-channel access, making it a unique solution for small cities.')

state_table = doc.add_table(rows=4, cols=5)
state_table.style = 'Table Grid'
state_table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr_cells_state = state_table.rows[0].cells
hdr_texts_state = ['Sr. No', 'Patent / Reference ID', 'Abstract / Description', 'Research Gap', 'Novelty']
for i, text in enumerate(hdr_texts_state):
    p = hdr_cells_state[i].paragraphs[0]
    p.add_run(text).bold = True
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

table_data = [
    ('1', 'US20180317209A1', 'Transit data platform for urban areas.', 'High-connectivity focus; no inclusivity.', 'Low-bandwidth operation with SMS/IVR access.'),
    ('2', 'CN111947982A', 'Scheduling system for public transport.', 'No crowd management or broad access.', 'AI-driven crowd predictions and multi-channel access.'),
    ('3', 'IN202221002153A', 'IoT-based tracking for limited areas.', 'Hardware-dependent; not scalable.', 'Software-only, scalable with management tools.')
]
for i, row_data in enumerate(table_data):
    row_cells = state_table.rows[i+1].cells
    for j, cell_text in enumerate(row_data):
        row_cells[j].text = cell_text

# D. DETAILED DESCRIPTION
add_section_heading(doc, 'D. Detailed Description:')
doc.add_paragraph('The AI-Powered Real-Time Public Transport Tracking and Crowd Density Management System is a next-generation software platform designed to address the unique challenges of public transportation in small urban centers. It aggregates real-time data from transport operators, such as driver updates or open transit feeds, into a cloud-based system. Advanced AI algorithms analyze this data to provide accurate arrival predictions and crowd density estimates, enabling commuters to plan trips effectively and avoid overcrowded vehicles.')
doc.add_paragraph('The platform is accessible through a user-friendly website, SMS updates, or IVR services, ensuring inclusivity for all residents, including those without smartphones. The web interface offers interactive maps, real-time alerts, and trip planning tools, while SMS and IVR provide simplified updates for non-digital users. For authorities, a management dashboard visualizes fleet performance, passenger distribution, and operational trends, facilitating data-driven decisions to optimize resources.')
doc.add_paragraph('Built with modern web technologies, the platform ensures seamless performance and scalability. Its lightweight design, using efficient data protocols, supports operation in low-connectivity environments, making it ideal for small cities. Security measures, including encrypted data transmission and compliance with privacy standards, protect user information, ensuring trust and reliability.')

# Software Components
doc.add_paragraph().add_run('Software Components:').bold = True
software = [
    ('1. Data Integration Module', [
        ('Component Name', 'Transit Data Aggregator'),
        ('Function', 'Collects real-time updates from transport sources.'),
        ('Features', 'Supports multiple data formats, operates in low-bandwidth conditions.')
    ]),
    ('2. Crowd Analytics Module', [
        ('Component Name', 'AI Crowd Predictor'),
        ('Function', 'Estimates passenger loads using user inputs and patterns.'),
        ('Technology', 'Machine learning-based analytics.'),
        ('Features', 'Hardware-free, real-time updates, high accuracy.')
    ]),
    ('3. Communication Module', [
        ('Component Name', 'Multi-Channel Notifier'),
        ('Function', 'Delivers updates via SMS or IVR.'),
        ('Protocol', 'Secure, lightweight messaging protocols.'),
        ('Features', 'Offline functionality, accessible to non-digital users.')
    ]),
    ('4. User Interface', [
        ('Component Name', 'Transit Web Interface'),
        ('Function', 'Provides maps, alerts, and trip planning tools.'),
        ('Platform', 'Browser-based, mobile-friendly.'),
        ('Features', 'Intuitive design, accessibility for diverse users.')
    ]),
    ('5. Analytics Backend', [
        ('Component Name', 'AI Predictive Engine'),
        ('Function', 'Processes data for arrival and crowd predictions.'),
        ('Features', 'Scalable, continuously learns from new data.')
    ])
]

for sw_title, sw_details in software:
    doc.add_paragraph().add_run(sw_title).bold = True
    for key, value in sw_details:
        add_indented_detail(doc, key, value)
    doc.add_paragraph()

# E. RESULTS AND ADVANTAGES
add_section_heading(doc, 'E. Results and Advantages: Comparison with existing solutions')
doc.add_paragraph('Testing in simulated small-city environments demonstrated the platform’s ability to deliver highly accurate arrival predictions, significantly reducing commuter wait times. The system improved fleet utilization by enabling data-driven scheduling, ensuring balanced passenger loads and minimizing overcrowding. Its inclusive design, with SMS and IVR access, received strong positive feedback, highlighting its accessibility and user-friendliness.')
doc.add_paragraph('Compared to existing solutions, the platform offers real-time transparency, enabling effective trip planning. Its crowd density predictions enhance safety by helping users avoid packed vehicles. The software-only approach reduces costs, making it feasible for small municipalities, while its low-bandwidth operation ensures reliability in areas with poor connectivity.')
doc.add_paragraph('The system’s superiority lies in its tailored approach for small cities, addressing gaps in inclusivity and connectivity. Early testing showed improved user satisfaction and operational efficiency, with potential for significant environmental benefits through reduced private vehicle use.')

# F. EXPANSION
add_section_heading(doc, 'F. Expansion: Potential extensions of the invention')
doc.add_paragraph('The platform is designed for scalability, with the potential to integrate digital payment systems for seamless ticketing, connect with electric vehicle fleets for eco-friendly monitoring, or link to smart city platforms for broader urban planning. It can also be adapted for private transport networks, such as school or corporate shuttles, expanding its applicability.')
doc.add_paragraph('Future enhancements could include predictive maintenance alerts or integration with real-time traffic management systems. These expansions would further enhance the platform’s utility, making it a cornerstone of smart urban mobility in small cities.')

# G. WORKING PROTOTYPE/ FORMULATION/ DESIGN/COMPOSITION
add_section_heading(doc, 'G. Working Prototype/ Formulation/ Design/Composition:')
doc.add_paragraph('The prototype is a fully functional web platform with a backend for data processing, a front-end interface for commuters, and a dashboard for authorities. Data flows seamlessly from transit sources to the analytics core, delivering real-time updates with high accuracy. The system uses efficient protocols to ensure reliability in low-connectivity settings.')
doc.add_paragraph('User-tested interfaces ensure accessibility for diverse populations, with iterative refinements based on feedback. The platform’s scalable architecture supports continuous updates, incorporating new data to improve predictions and enhance user experience over time.')

# G. EXISTING DATA
add_section_heading(doc, 'G. Existing Data: Supporting evidence for the invention')
doc.add_paragraph('Testing demonstrated high accuracy in arrival predictions and effective crowd density estimation, with the platform handling multiple data streams without performance issues. Compared to existing tools, it excels in low-connectivity environments, ensuring reliable access for all users.')
doc.add_paragraph('User feedback confirmed the system’s ease of use and inclusivity, with strong performance metrics in simulated environments. These results validate the platform’s readiness for deployment and its potential to transform public transportation in small cities.')

# 4. USE AND DISCLOSURE
p_use = doc.add_paragraph()
p_use.paragraph_format.space_before = Pt(12)
p_use.add_run('4. Use and Disclosure (Important): ').bold = True
p_use.add_run('Please answer the following questions:')

disclosure_table = doc.add_table(rows=6, cols=3)
disclosure_table.style = 'Table Grid'
disclosure_table.autofit = False
disclosure_table.columns[0].width = Inches(4.5)
disclosure_table.columns[1].width = Inches(1.0)
disclosure_table.columns[2].width = Inches(1.0)

hdr_cells = disclosure_table.rows[0].cells
hdr_cells[1].text = 'YES'
hdr_cells[2].text = 'NO'
hdr_cells[1].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
hdr_cells[2].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

questions = [
    'Have you described or shown your invention/design to anyone or in any conference?',
    'Have you made any attempts to commercialize your invention (e.g., approached companies)?',
    'Has your invention been described in any publication or media, such as the Internet?',
    'Do you have any collaboration with other institutes/organizations? Provide details.',
    'Name of regulatory body or approvals required, if any.'
]

for i, q in enumerate(questions):
    row_cells = disclosure_table.rows[i+1].cells
    row_cells[0].text = q
    row_cells[1].text = '(  )'
    row_cells[2].text = '( ü )'
    row_cells[1].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    row_cells[2].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
doc.add_paragraph()

# Sections 5-10
use_questions = [
    ('5. Provide links and dates for any public disclosures (e.g., papers, videos) before sharing with us.', 'Not applicable.'),
    ('6. Provide terms of any MOU for collaborations with other entities.', 'Not applicable.'),
    ('7. Potential Chances of Commercialization.', 'The platform has strong potential for adoption by municipal transport authorities and technology providers, particularly in emerging markets with growing urban populations.'),
    ('8. List of companies which can be contacted for commercialization with website links.', 'Not applicable.'),
    ('9. Any basic patent requiring royalty payments.', 'Not applicable.'),
    ('10. Filing Options: Indicate the level of work for provisional/complete/PCT filings.', 'Complete Patent')
]

for q, a in use_questions:
    p_q = doc.add_paragraph()
    p_q.add_run(q).bold = True
    p_a = doc.add_paragraph(a)
    p_a.paragraph_format.left_indent = Inches(0.5)

# 11. KEYWORDS
p_kw = doc.add_paragraph()
p_kw.paragraph_format.space_before = Pt(12)
p_kw.add_run('11. Keywords: ').bold = True
p_kw.add_run('Keywords for searching the invention.')

keywords = [
    'AI Public Transport Tracker',
    'Real-Time Transit System',
    'Crowd Density Management',
    'Small City Mobility Solution',
    'Inclusive Transport Platform',
    'AI Arrival Predictions',
    'Web-Based Transit Tool',
    'Low-Bandwidth Transport System',
    'Sustainable Urban Mobility',
    'Transport Management Dashboard'
]
for kw in keywords:
    doc.add_paragraph(kw, style='List Bullet')

# NO OBJECTION CERTIFICATE
add_section_heading(doc, 'No Objection Certificate', align=WD_PARAGRAPH_ALIGNMENT.CENTER, space_before=24)
doc.add_paragraph('This is to certify that [University/Organization Name] or its associates have no objection to Lovely Professional University filing an IPR (Patent/Copyright/Design/etc.) entitled "AI-Powered Real-Time Public Transport Tracking and Crowd Density Management System for Small Cities" including the name(s) of inventors who are students/employees of our institution.')
doc.add_paragraph('Further, [University/Organization Name] will not provide financial assistance for this IPR nor raise objections regarding its filing, commercialization, or claim any rights to the invention at any stage.')
doc.add_paragraph('\n\n(Authorised Signatory)')

# --- 3. Save the Document ---
try:
    file_name = 'AI_Public_Transport_Tracking_System_Annexure3b.docx'
    doc.save(file_name)
    print(f'DOCX file created successfully! Check the current directory for {file_name}')
except Exception as e:
    print(f"Error: Could not save the file. Make sure '{file_name}' is not open.")
    print(f"Details: {e}")